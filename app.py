import streamlit as st
from docx import Document
import io
import os
import requests


GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_API_KEY = "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"  

def edit_word_file(file, prompt):
    doc = Document(file)
    edited_text = []

    
    for para in doc.paragraphs:
        edited_text.append(para.text)

    full_text = "\n".join(edited_text)

    
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "llama3-8b-8192",  
        "messages": [
            {"role": "system", "content": "You are a helpful editor. Edit the text as per user instructions."},
            {"role": "user", "content": f"Here is the text:\n{full_text}\n\nPlease edit it according to this instruction:\n{prompt}\n\nReturn the full edited text."}
        ]
    }
    response = requests.post(GROQ_API_URL, headers=headers, json=data)
    response.raise_for_status()
    edited_text = response.json()["choices"][0]["message"]["content"]

    
    edited_doc = Document()
    for line in edited_text.split("\n"):
        edited_doc.add_paragraph(line)

    
    edited_stream = io.BytesIO()
    edited_doc.save(edited_stream)
    edited_stream.seek(0)
    return edited_stream

def main():
    st.title("File Editor Agent")

    uploaded_file = st.file_uploader("Upload a Word file (.docx)", type="docx")
    prompt = st.text_area("Enter editing instructions", "Replace 'pwer' with 'power'")

    if uploaded_file and prompt:
        if st.button("Edit File"):
            with st.spinner("Editing..."):
                edited_stream = edit_word_file(uploaded_file, prompt)
                st.success("Editing complete!")
                st.download_button(
                    label="Download Edited File",
                    data=edited_stream,
                    file_name="edited_file.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
